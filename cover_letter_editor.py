from docx import Document
from docx.shared import Pt

def edit(company_name='google', job_name = "computer scientist"):
    # Load template and create new document
    doc_path = "your document path here"
    template = Document(doc_path)
    new_cover = Document() 

    # Iterate over the paragraphs in the template
    for paragraph in template.paragraphs:
        # Replace the placeholder text with the actual values
        t = paragraph.text.replace("COMPANYNAME", company_name)
        t = t.replace("JOBROLE", job_name)

        # Add a new paragraph to the new_cover document
        para = new_cover.add_paragraph().add_run(t)

        # Set the font size and font name
        para.font.size = Pt(12)
        para.font.name = "Times New Roman"

    # saving the modified document
    new_cover.save("assets\cover_letter.docx")

if __name__  == "__main__":
    edit()
